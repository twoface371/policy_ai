"""
analyzer.py — 문서 법령 저촉 분석기 (단독 검증용)
================================================================================
나중에 v10에 합칠 '분석 엔진'만 따로 뗀 것. 웹서버·DB 없음.

기능:
  1. 파일에서 텍스트 추출 (docx / pdf / hwp / hwpx / txt 자동 감지)  ← LLM 0
  2. 법령 인용 탐지 (정규식)  ← LLM 0
  3. (다음 단계) 현행법 대조 — 여기선 뼈대만, 실제 대조는 법제처 API 연결 후

사용:
  python analyzer.py 파일경로              # 파일 분석
  python analyzer.py --text "○○법 제5조"   # 텍스트 직접 분석
  python analyzer.py --selftest            # 내장 테스트
================================================================================
"""
import re
import sys
import os
import subprocess
import zipfile
from typing import List, Dict


# ============================================================
# 1. 텍스트 추출 (포맷 자동 감지)
# ============================================================
def extract_text(path: str) -> str:
    """확장자로 포맷을 판단해 텍스트를 뽑는다. 실패 시 예외 대신 빈 문자열."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return _from_docx(path)
    if ext == ".pdf":
        return _from_pdf(path)
    if ext == ".hwp":
        return _from_hwp(path)
    if ext == ".hwpx":
        return _from_hwpx(path)
    if ext in (".txt", ".md"):
        with open(path, encoding="utf-8", errors="replace") as f:
            return f.read()
    raise ValueError(f"지원하지 않는 형식: {ext} (docx/pdf/hwp/hwpx/txt)")


def _from_docx(path: str) -> str:
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    # 표 안의 텍스트도 수집 (법령이 표에 들어있는 경우 많음)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _from_pdf(path: str) -> str:
    import pdfplumber
    out = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n".join(out)


def _hwp5txt_cmd() -> str:
    """hwp5txt 실행 파일 경로.

    pyhwp는 실행 파일을 파이썬 옆(가상환경의 bin)에 깐다. 서버를
    `.venv/bin/python app.py`로 띄우면 그 디렉터리가 PATH에 없어서
    이름만으로는 못 찾는다 — 그래서 hwp가 통째로 실패했다.
    지금 도는 파이썬 옆을 먼저 보고, 없으면 PATH에 맡긴다.
    """
    cand = os.path.join(os.path.dirname(sys.executable), "hwp5txt")
    return cand if os.path.exists(cand) else "hwp5txt"


def _from_hwp(path: str) -> str:
    """한글 5.0 (.hwp) — pyhwp의 hwp5txt 커맨드 사용."""
    try:
        r = subprocess.run([_hwp5txt_cmd(), path], capture_output=True, timeout=60)
        if r.returncode == 0:
            return r.stdout.decode("utf-8", errors="replace")
        # 커맨드 실패 시 파이썬 API 시도
        return _from_hwp_api(path)
    except FileNotFoundError:
        return _from_hwp_api(path)


def _from_hwp_api(path: str) -> str:
    """hwp5txt 커맨드가 없을 때 파이썬 API로.

    transform_hwp5_to_text는 바이트를 쓴다. StringIO를 주면
    'string argument expected, got bytes'로 죽는다.
    """
    try:
        from hwp5.xmlmodel import Hwp5File
        from hwp5.hwp5txt import TextTransform
        import io
        buf = io.BytesIO()
        hwp = Hwp5File(path)
        TextTransform().transform_hwp5_to_text(hwp, buf)
        return buf.getvalue().decode("utf-8", errors="replace")
    except Exception as e:
        raise RuntimeError(f"hwp 추출 실패: {e}. hwp5txt 설치 필요 (pip install pyhwp)")


def _from_hwpx(path: str) -> str:
    """한글 신형 (.hwpx) — zip 기반. section*.xml에서 텍스트 추출."""
    texts = []
    with zipfile.ZipFile(path) as z:
        sections = sorted(n for n in z.namelist()
                          if re.match(r"Contents/section\d+\.xml", n))
        for name in sections:
            xml = z.read(name).decode("utf-8", errors="replace")
            # <hp:t>...</hp:t> 안의 텍스트가 본문
            for m in re.finditer(r"<hp:t>(.*?)</hp:t>", xml, re.DOTALL):
                t = re.sub(r"<[^>]+>", "", m.group(1))
                if t.strip():
                    texts.append(t)
    return "\n".join(texts)


# ============================================================
# 2. 법령 인용 탐지 (정규식) — LLM 0
# ============================================================
# 실제 공공 법률 문서를 기준으로 4가지 패턴을 처리한다:
#   A. 따옴표/괄호로 감싼 법령명:  '국가를 당사자로 하는 계약에 관한 법률' 제27조
#                                「소프트웨어 진흥법」 제51조제7항
#   B. 따옴표 없는 법령명 + 조문:   중소기업기본법 제2조
#   C. 동법 / 같은 법 (이어받기):   동법 시행령 제76조  (앞에 나온 법을 가리킴)
#   D. 조 나열·범위 (이어받기):     제3조 및 제5조 / 제3조부터 제7조까지
#
# B는 이름의 시작을 확정할 수 없다(조사가 이름 안에도 밖에도 나온다). 그래서
# 하나로 정하지 않고 후보를 짧은 것부터 여러 개 내보내고, checker가 법제처
# 조회로 고른다. 실측 재현율은 첫 후보 84% / 후보 포함 100%.

_QUOTE = r"['\u2018\u2019\u201c\u201d\u300c\u300d\u300e\u300f]"  # ' ' " " 「 」 『 』
# 긴 것을 앞에 둔다. '법'이 먼저면 '법률'에서 '법'만 물고 '률'이 남는다.
_TAIL = r"(?:법률|법|시행규칙|시행령|규정|규칙|지침|예규|조건|기준)"

_LAW_QUOTED = re.compile(
    _QUOTE + r"\s*([가-힣A-Za-z0-9·ㆍ\s]+?" + _TAIL + r")\s*" + _QUOTE +
    r"\s*제(\d+)조(?:의(\d+))?(?:\s*제(\d+)항)?(?:\s*제(\d+)호)?")

_LAW_BARE = re.compile(
    r"제(\d+)조(?:의(\d+))?(?:\s*제(\d+)항)?(?:\s*제(\d+)호)?")

# E. 조문 없이 법령명만 — '「개인정보 보호법」에 따라', '전자서명법에 따른'
#
# 위 A~D가 전부 '제N조'를 요구해서, 법 전체를 가리키는 문장이 통째로 탐지에서
# 빠졌다. 실제 공문에 흔한 형태다.
#
# 이 패턴은 앞의 것들보다 훨씬 헐겁다. '방법·위법·현행법'처럼 법으로 끝나는
# 보통 낱말이 다 걸리기 때문이다. 그래서 두 겹으로 막는다.
#   1) 아래 _NOT_LAW 로 뻔한 것을 먼저 걷어낸다.
#   2) 남은 것은 법제처에 실제로 있는 이름일 때만 살린다(checker 가 처리).
#      조가 없으니 '없는 법령'이라고 지적할 근거도 약하다 — 조용히 버린다.
_LAW_ONLY_QUOTED = re.compile(
    _QUOTE + r"\s*([가-힣A-Za-z0-9·ㆍ\s]+?" + _TAIL + r")\s*" + _QUOTE)

# 전문 참조에서는 꼬리를 좁힌다. 규정·규칙·지침·예규·조건·기준은 보통명사로도
# 쓰여서('세부 기준', '무처리 규정', '이 규정'), 조문이 없으면 법령인지 문장의
# 일부인지 가릴 수 없다. 실측에서 오탐의 대부분이 이 여섯이었다.
# 조문이 붙은 인용은 '제N조'가 법령임을 보증하므로 기존 _TAIL 을 그대로 쓴다.
_TAIL_STRICT = r"(?:법률|법|시행규칙|시행령)"

# 법령명은 한 줄 안에서만 이어 붙인다(_LAW_TAIL 과 같은 이유).
# 뒤에 조사·구두점·줄끝이 와야 한다 — '개인정보 보호법률안' 같은 말의 중간을
# 물지 않게 한다.
_LAW_ONLY_BARE = re.compile(
    r"([가-힣·ㆍ]+(?:[ \t][가-힣·ㆍ]+){0,7}?[ \t]?" + _TAIL_STRICT + r")"
    r"(?=[\s,.·]|을|를|은|는|이|가|에|의|와|과|상|및|또는|$)")

# 관형사로 시작하면 법령명이 아니다 — '그 방법', '이 시행령', '해당 법률'.
_DETERMINER = re.compile(r"^(?:그|이|저|본|해당|동|각|위|아래|같은|앞|뒤|당해)\s")

# 법으로 끝나지만 법령명이 아닌 말. 이것을 인용으로 올리면 검사 결과가
# 노이즈로 덮인다.
_NOT_LAW = {
    "방법", "위법", "적법", "불법", "합법", "준법", "탈법", "편법", "입법",
    "사법", "행정법", "현행법", "국내법", "국제법", "실정법", "성문법",
    "관습법", "특별법", "일반법", "모법", "상위법", "하위법", "개별법",
    "관련법", "관계법", "이 법", "그 법", "본 법", "해당 법", "같은 법",
    "동법", "법률", "법", "시행령", "시행규칙", "규정", "규칙", "지침",
    "예규", "조건", "기준", "관련 법", "관계 법", "관련 법령", "관계 법령",
}

# 법령명 바로 뒤에 조문이 오면 그것은 조문 인용이다(전문 참조가 아니다).
_FOLLOWED_BY_ART = re.compile(r"\s*제\s*\d+\s*조")

# 조문 조각에서 시작하는 것을 거른다 — '제3조 및 제5조의 규정' 에서
# '조의 규정' 같은 토막이 잡힌다.
_ART_FRAGMENT = re.compile(r"^(?:제?\s*\d+\s*)?[조항호목]")


def _plausible_law_name(law: str) -> bool:
    """전문 참조로 올릴 만한 이름인가. 값싼 판정만 한다.

    최종 판단은 법제처 조회가 한다(checker). 여기서는 조회를 낭비하지 않을
    만큼만 걸러 낸다.
    """
    if law in _NOT_LAW or len(law.replace(" ", "")) < 2:
        return False
    if _ART_FRAGMENT.match(law) or any(ch.isdigit() for ch in law):
        return False
    if _DETERMINER.match(law):
        return False
    # 앞이 잘려 나온 조각 — '행 · 지침'처럼 한 음절로 시작하는 두 낱말짜리.
    # 진짜 한 음절 법령명(민법·형법)은 낱말이 하나라 여기 걸리지 않는다.
    tokens = law.split()
    return not (len(tokens) > 1 and len(tokens[0]) == 1)

# 법령명은 한 줄 안에서만 이어 붙인다. \s로 두면 개행을 넘어가서
# 바로 윗줄(소제목 등)의 단어까지 법령명으로 끌어온다.
#
# 꼬리(법률/법/…) 앞의 [ \t]? 가 핵심이다. 이게 없으면 꼬리가 '독립된 한 단어'인
# 이름을 못 잡는다 — 반복 그룹의 [가-힣]+ 가 '법률'을 통째로 삼켜 꼬리에 남길
# 글자가 없어지기 때문이다. '…에 관한 법률'은 국내 법률명에서 가장 흔한 형태라
# 이것 하나로 긴 법률이 통째로 탐지에서 빠졌다.
_LAW_TAIL = re.compile(
    r"([가-힣·ㆍ]+(?:[ \t][가-힣·ㆍ]+)*?[ \t]?" + _TAIL + r")[ \t]*$")

_DONGBEOP = re.compile(
    r"(동법|같은\s?법)\s*(시행령|시행규칙)?\s*"
    r"제(\d+)조(?:의(\d+))?(?:\s*제(\d+)항)?(?:\s*제(\d+)호)?")

_JOSA_END = re.compile(r"(은|는|이|가|을|를|에|의|와|과|로|으로|및|또는)$")

# '하는/되는'의 '는'은 조사가 아니라 관형형 어미다. 뒤 명사를 꾸미므로
# 언제나 이름 안쪽인데 _JOSA_END의 '는'에 걸려 경계로 오인된다.
# '국가를 당사자로 하는 계약에 관한 법률'이 '계약에 관한 법률'로 잘리던 원인.
_ADNOMINAL = re.compile(r"(하는|되는|받는|드는|오는|가는|있는|없는|같은)$")

# 법령명을 거슬러 찾을 때 볼 앞 구간 길이. 국내 법령명 중 가장 긴 것이
# 60자 안쪽이라 넉넉하다.
_LOOKBACK = 200

# 꼬리만 남은 이름('이 법 제3조' → '법'). 그 문서 자신을 가리키는 자기참조라
# 어느 법인지 알 수 없고, 검색해 봐야 '확인 필요'만 쌓인다. 실제 법령 본문
# 116만 자에서 인용 1,591건 중 93건이 이 형태였다.
_BARE_TAIL = re.compile(r"^(?:이|본|당해|해당|위|그)?\s*" + _TAIL + r"$")

# 조사로 끝나도 법령명이 계속되는 연결어. '공공기관의 정보공개에 관한 법률'을
# '정보공개에'의 '에'에서 끊으면 '관한 법률'만 남아 검색이 안 된다.
# 국내 법률명에 아주 흔한 형태라 이 예외가 없으면 긴 법률이 통째로 누락된다.
_NAME_CONNECTOR = ("관한", "대한", "위한", "따른", "의한")

# 법령명 앞에 붙는 접속부사. 법령명에는 절대 들어가지 않는데 조사로 끝나지
# 않아 조사 경계 처리로는 안 걸린다. ('또한 국가재정법' → '또한'까지 딸려옴)
_LEAD_STOP = ("또한", "다만", "그리고", "그러나", "아울러", "특히", "한편",
              "따라서", "이때", "이에", "즉", "만약", "만일", "우선", "먼저",
              "다음", "이번", "향후", "현재")

# 꼬리로 끝나는지 판별 — '및' 앞이 법령명의 끝인지 보는 데 쓴다
_ENDS_TAIL = re.compile(_TAIL + r"$")

# 조 나열·범위에서 법령명을 이어받기 위한 연결 구간.
#   '제3조 및 제5조'  '제3조ㆍ제5조'  '제3조, 제5조'  '제3조부터 제7조까지'
# 이 사이에 다른 말이 끼면 이어받지 않는다 — 앞 법령이 계속된다는 근거가 없다.
_CONNECTOR_GAP = re.compile(
    r"^\s*(?:및|와|과|·|ㆍ|,|~|부터|내지|또는|그리고)?\s*$")

# 줄바꿈으로 잘린 법령명을 잇는다. 다만 무조건 이으면 윗줄 소제목까지
# 끌어오므로('제1장 총칙' + '법률'), 윗줄이 '…에 관한'처럼 이어질 수밖에 없는
# 연결어로 끝날 때만 넘어간다. 긴 법률명이 잘리는 자리가 대부분 여기다.
_WRAP_JOIN = re.compile(r"(?:" + "|".join(_NAME_CONNECTOR) + r")[ \t]*\n[ \t]*")


def _build_article(nums) -> str:
    # 조문 없이 법령명만 인용한 경우. 빈 문자열이라야 화면·리포트의 조문 칸이
    # 비어 나온다 — 그러지 않으면 '제None조'가 찍힌다.
    if not nums or nums[0] is None:
        return ""
    a = f"제{nums[0]}조"
    if nums[1]:
        a += f"의{nums[1]}"
    if nums[2]:
        a += f"제{nums[2]}항"
    if nums[3]:
        a += f"제{nums[3]}호"
    return a


def _art_key(nums):
    """_build_article과 같은 슬라이스에서 조 번호만 숫자로 뽑는다.
    표시용 문자열만 남기면 현행 조문과 대조할 좌표가 사라진다.

    조문 없이 법령명만 인용한 경우는 (0, 0)이다 — 대조할 좌표가 아예 없다.
    """
    if not nums or nums[0] is None:
        return 0, 0
    return int(nums[0]), int(nums[1] or 0)


def _base_law(name: str) -> str:
    """법령명에서 시행령/시행규칙 꼬리를 떼어 '원 법률'을 구함.
    (동법이 가리키는 것은 시행령이 아니라 그 원 법률)"""
    return re.sub(r"\s*(?:시행령|시행규칙)\s*$", "", name).strip()


def _ok_name(nm: str) -> bool:
    """법령명으로 성립하지 않는 후보를 거른다.

    '법'이 독립 토큰으로 끝나면 자기참조다 — '이 법 제3조', '그 경우 법 제5조'.
    저장된 실제 법령명 80개 중 이렇게 끝나는 것은 하나도 없고, '법률'로 끝나는
    것은 21개다('…에 관한 법률'). 그래서 '법'만 거르고 '법률'은 남긴다.
    """
    if _BARE_TAIL.match(nm):
        return False
    toks = nm.split()
    return not (len(toks) > 1 and toks[-1] == "법")


def _law_name_candidates(text: str, pos: int, limit: int = 3) -> List[str]:
    """pos(‘제N조’ 시작) 직전 구간에서 법령명 후보를 짧은 것부터 만든다.

    하나로 확정하지 않는 이유가 있다. 조사 경계로는 이름의 시작을 결정할 수
    없다 — '공공기관의 정보공개에 관한 법률'의 '의'는 이름 안이고 '이 사업은
    …법 제3조'의 '은'은 문장 경계인데, 겉모습이 같다. 실제로 실측하면 이름의
    16%가 앞이 잘렸고, 잘린 이름이 **다른 법에 정확히 맞아버리는** 경우까지
    나왔다('보호 및 지원에 관한 법률' → 북한이탈주민…). 조용히 틀린 법을
    대조하느니 후보를 넘겨 법제처가 고르게 한다(checker.check_name).

    첫 후보는 지금까지 쓰던 값(마지막 조사 경계)이다. 이것이 실측 84%에서
    바로 맞으므로 대부분은 추가 조회 없이 끝난다.
    """
    # 조문 표기 바로 앞 구간만 본다. text[:pos]를 통째로 넘기면 인용마다
    # 문서 전체를 복사하고 정규식으로 훑어 O(n²)가 된다(100만 자 문서에서
    # 수십 초). 법령명은 가장 긴 것도 60자를 넘지 않아 이 창이면 충분하다.
    before = text[max(0, pos - _LOOKBACK):pos].rstrip()
    # '…에 관한\n법률'처럼 줄바꿈으로 잘린 이름을 먼저 붙인다
    before = _WRAP_JOIN.sub(lambda mo: mo.group(0).split("\n")[0].strip() + " ",
                            before)
    m = _LAW_TAIL.search(before)
    if not m:
        return []
    tokens = m.group(1).split()
    cuts = []
    for i, tk in enumerate(tokens[:-1]):
        if not _JOSA_END.search(tk) or _ADNOMINAL.search(tk):
            continue
        # 뒤따르는 말이 '관한/대한/…'이면 조사가 아니라 법령명 내부다
        if tokens[i + 1] in _NAME_CONNECTOR:
            continue
        # '및/또는'은 법령명 안에도 흔하다('정보통신망 이용촉진 및 정보보호
        # 등에 관한 법률'). 바로 앞이 법 꼬리로 끝날 때만 문장 접속으로 보고
        # 자른다 — '개인정보 보호법 및 …'은 잘라야 맞고, '이용촉진 및 …'은
        # 자르면 이름이 토막 난다.
        mo = re.search(r"(?:및|또는)$", tk)
        if mo:
            head = tk[:mo.start()] or (tokens[i - 1] if i else "")
            if not _ENDS_TAIL.search(head):
                continue
        cuts.append(i + 1)

    # 짧은 것(마지막 경계) → 점점 긴 것 → 통째로
    starts = ([cuts[-1]] if cuts else []) + list(reversed(cuts[:-1]))
    if 0 not in starts:
        starts.append(0)

    out: List[str] = []
    for s in starts:
        # 접속부사가 앞에 붙어 있으면 떼어 낸다
        while s < len(tokens) - 1 and tokens[s] in _LEAD_STOP:
            s += 1
        nm = " ".join(tokens[s:])
        if nm and _ok_name(nm) and nm not in out:
            out.append(nm)
        if len(out) >= limit:
            break
    return out


def find_citations(text: str) -> List[Dict]:
    """법령 인용을 찾아 [{법령, 조문, 문맥}] 리스트로 반환.
    따옴표 인용 / 일반 인용 / 동법 이어받기를 모두 처리하고,
    겹치는 구간은 긴 것(따옴표 인용)을 우선한다."""
    events = []
    for m in _LAW_QUOTED.finditer(text):
        events.append((m.start(), m.end(), "quoted", m))
    for m in _DONGBEOP.finditer(text):
        events.append((m.start(), m.end(), "dong", m))
    for m in _LAW_BARE.finditer(text):
        events.append((m.start(), m.end(), "bare", m))
    # 조문 없는 법령명은 맨 뒤에 둔다. 정렬이 (위치, 길이)라 같은 자리에서
    # 조문까지 붙은 긴 인용이 먼저 잡히고, 이것은 남은 자리에만 들어간다 —
    # '개인정보 보호법 제15조'가 '개인정보 보호법'으로 축소되면 안 된다.
    for m in _LAW_ONLY_QUOTED.finditer(text):
        events.append((m.start(), m.end(), "lawq", m))
    for m in _LAW_ONLY_BARE.finditer(text):
        events.append((m.start(), m.end(), "lawb", m))
    # 3순위 키: 조문이 붙은 것(0)을 조문 없는 것(1)보다 앞세운다. 길이만으로는
    # '「소프트웨어 진흥법」'(따옴표 포함)이 '소프트웨어 진흥법 제51조'보다
    # 길어지는 자리가 생긴다.
    _rank = {"quoted": 0, "dong": 0, "bare": 0, "lawq": 1, "lawb": 1}
    events.sort(key=lambda x: (x[0], _rank[x[2]], -(x[1] - x[0])))

    claimed = []
    out = []
    last_law = None

    def overlaps(s, e):
        return any(not (e <= cs or s >= ce) for cs, ce in claimed)

    for s, e, kind, m in events:
        if overlaps(s, e):
            continue
        if kind == "quoted":
            # 따옴표가 경계를 확정해 주므로 후보를 나눌 이유가 없다
            law = re.sub(r"\s+", " ", m.group(1)).strip()
            cands = [law]
            nums = m.groups()[1:5]
            last_law = law
        elif kind == "dong":
            base = _base_law(last_law) if last_law else "(앞 법령 불명)"
            suffix = m.group(2) or ""
            law = f"{base} {suffix}".strip() if suffix else base
            cands = [law]
            nums = m.groups()[2:6]
        elif kind in ("lawq", "lawb"):
            # 조문 없이 법령명만.
            if kind == "lawq":
                # 따옴표가 양쪽 경계를 확정해 준다
                law = re.sub(r"\s+", " ", m.group(1)).strip()
            else:
                # 따옴표가 없으면 왼쪽 경계를 정규식이 못 정한다. 조문 인용과
                # 같은 절단 로직을 태워 앞 문장을 떼어 낸다 — 그러지 않으면
                # '본 사업은 전자정부법'처럼 문장째로 잡힌다.
                cands = _law_name_candidates(text, e)
                if not cands:
                    continue
                law = cands[0]
            if not _plausible_law_name(law):
                continue
            # 바로 뒤에 조문이 붙어 있으면 그것은 조문 인용이다. 조문 쪽 이벤트가
            # '제N조'만 차지해 겹침 검사에 안 걸리므로 여기서 걸러야 한다.
            if _FOLLOWED_BY_ART.match(text[e:e + 12]):
                continue
            cands = [law]
            last_law = law
            nums = (None, None, None, None)
        else:  # bare
            cands = _law_name_candidates(text, s)
            if not cands and last_law and claimed \
                    and _CONNECTOR_GAP.match(text[claimed[-1][1]:s]):
                # '제3조 및 제5조', '제3조부터 제7조까지' — 둘째 조 앞에는
                # 법령명이 없다. 사이에 연결어밖에 없으면 앞 법령이 계속된다.
                cands = [last_law]
            if not cands:
                continue
            law = cands[0]
            last_law = law
            nums = m.groups()
        claimed.append((s, e))
        art_no, art_branch = _art_key(nums)
        ctx_s = max(0, s - 40)
        ctx_e = min(len(text), e + 40)
        out.append({"법령": law, "법령후보": cands,
                    "조문": _build_article(nums),
                    "조번호": art_no, "조가지번호": art_branch,
                    # 전문 참조 — 조를 대조할 대상이 없다. checker 가 이 값을
                    # 보고 조 판정을 건너뛰고, 현행에 없는 이름이면 버린다.
                    "전문참조": kind in ("lawq", "lawb"),
                    "문맥": text[ctx_s:ctx_e].replace("\n", " ").strip()})
    return out


# ============================================================
# 실행 / 테스트
# ============================================================
def analyze_text(text: str) -> Dict:
    cites = find_citations(text)
    return {"글자수": len(text), "인용수": len(cites), "인용": cites}


def _selftest():
    sample = """개인정보 처리 방침
본 방침은 개인정보 보호법 제15조에 따라 개인정보를 수집·이용합니다.
정보주체는 국가정보화 기본법 제3조에 명시된 권리를 가집니다.
전자정부법 제2조제10호의 정의를 준용하며, 위반 시 제76조에 따라 처벌됩니다.
이 문서는 소프트웨어산업 진흥법 제20조를 근거로 작성되었습니다.
사회보장급여법 제5조의2제1항에 따른 조치를 취합니다.
일반 문장입니다. 법령 언급이 없는 줄도 있습니다."""
    r = analyze_text(sample)
    print(f"글자수 {r['글자수']} · 인용 {r['인용수']}건\n")
    # 조번호·조가지번호는 현행 조문 대조의 좌표라 표시 문자열과 함께 검증한다.
    expected = [
        ("개인정보 보호법", "제15조", 15, 0),
        ("국가정보화 기본법", "제3조", 3, 0),
        ("전자정부법", "제2조제10호", 2, 0),
        ("소프트웨어산업 진흥법", "제20조", 20, 0),
        ("사회보장급여법", "제5조의2제1항", 5, 2),
    ]
    got = [(c["법령"], c["조문"], c["조번호"], c["조가지번호"]) for c in r["인용"]]
    for law, art, no, br in got:
        print(f"  · {law}  {art}  → 조 {no}" + (f"의{br}" if br else ""))
    print()
    ok = got == expected
    if ok:
        print("✅ 셀프테스트 통과 — 인용 5건 정확히 탐지")
    else:
        print("❌ 불일치")
        print("  기대:", expected)
        print("  실제:", got)
    return ok


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(0)
    if sys.argv[1] == "--selftest":
        _selftest()
    elif sys.argv[1] == "--text":
        text = sys.argv[2] if len(sys.argv) > 2 else ""
        import json
        print(json.dumps(analyze_text(text), ensure_ascii=False, indent=2))
    else:
        path = sys.argv[1]
        print(f"[파일] {path}")
        text = extract_text(path)
        print(f"[추출] {len(text)}자\n")
        r = analyze_text(text)
        print(f"[인용] {r['인용수']}건")
        for c in r["인용"]:
            print(f"  · {c['법령']}  {c['조문']}")
            print(f"      문맥: …{c['문맥']}…")
